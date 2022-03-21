from docx import Document
import re

CITE_REGEX = re.compile('[a-zA-Z ]{2,40}[\d]{2,4}')

def unique_cards(cards):
    cs = []
    for card in cards:
        if card.tag in [_.tag for _ in cs] or card.text in [_.text for _ in cs]:
            print("Repeat Card -- skipping")
        else:
            cs.append(card)
    return cs

def extract_cards(filename): # tag: (citation, text)
    document = Document(filename)
    cards = []

    for i, paragraph in enumerate(document.paragraphs):
        tag = document.paragraphs[i-2].text
        cite = document.paragraphs[i-1].text
        text = paragraph.text
        reg_match = CITE_REGEX.findall(cite)
        if len(reg_match) > 0:
            if len(text) > 1000 and cite.strip() and tag.strip() and len(tag.strip().split()) > 2 \
                and cite.index(reg_match[0]) == 0: # TODO improve
                if (tag in [_[0] for _ in cards] or text in [_[2] for _ in cards]) or not tag.strip():
                    print("Repeat Card -- skipping")
                else:
                    cards.append((tag, cite, paragraph))
                
    return cards