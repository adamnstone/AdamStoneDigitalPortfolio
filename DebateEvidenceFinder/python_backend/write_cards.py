from docx import Document
from docx.shared import Pt, RGBColor

def add_card_to_document(card, document):
    head = document.add_heading().add_run(card.get_tag())
    head.bold = True
    head.font.color.rgb = RGBColor(0, 0, 0)
    head.font.size = Pt(13)
    cite = document.add_paragraph().add_run(card.citation)
    cite.font.size = Pt(8)
    para = document.add_paragraph()
    for run in card.text.runs:
        para_run = para.add_run(run.text)
        para_run.bold = run.bold
        para_run.underline = run.underline
        para_run.italic = run.italic
        para_run.font.highlight_color = run.font.highlight_color
        para_run.style = run.style

def write_cards(cards, filename):
    document = Document()

    for card in cards:
        add_card_to_document(card, document)

    document.save(filename)